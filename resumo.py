from docx import Document

# Cria um novo documento
doc = Document()
doc.add_heading('Resumo para Estudo - Demonstrações Contábeis e Custos no Setor Público', 0)

# Introdução às Demonstrações do Setor Público
doc.add_heading('Introdução às Demonstrações do Setor Público', level=1)
doc.add_paragraph('De acordo com a resolução CFC n. 1.133/08 (NBC T 16.6 (R1)), as demonstrações contábeis obrigatórias são:')
doc.add_paragraph('a. Balanço Patrimonial\nb. Demonstrações das Variações Patrimoniais\nc. Demonstração das Mutações do Patrimônio Líquido\n'
                   'd. Demonstração dos Fluxos de Caixa\ne. Balanço Orçamentário\nf. Balanço Financeiro\n'
                   'g. Notas explicativas, compreendendo a descrição sucinta das principais políticas contábeis e outras informações elucidativas\n'
                   'h. Informação comparativa com o período anterior')

# Balanço Orçamentário
doc.add_heading('Balanço Orçamentário', level=2)
doc.add_paragraph('O Balanço Orçamentário evidencia as receitas e as despesas orçamentárias, detalhadas em níveis relevantes de análise, '
                  'confrontando o orçamento inicial e as suas alterações com a execução, demonstrando o resultado orçamentário.')

# Balanço Financeiro
doc.add_heading('Balanço Financeiro', level=2)
doc.add_paragraph('O Balanço Financeiro evidencia as receitas e despesas orçamentárias, bem como os ingressos e dispêndios extraorçamentários, '
                  'conjugados com os saldos de caixa do exercício anterior e os que se transferem para o início do exercício seguinte.')

# Demonstração dos Fluxos de Caixa (DFC)
doc.add_heading('Demonstração dos Fluxos de Caixa (DFC)', level=2)
doc.add_paragraph('A DFC apresenta as entradas e saídas de caixa e as classifica em fluxos operacional, de investimentos e de financiamento.\n'
                  'Identifica as fontes de geração de caixa, os itens de consumo de caixa durante o período das demonstrações contábeis e o saldo de caixa na data das demonstrações contábeis.\n'
                  'Essa demonstração permite aos usuários saber como a entidade obteve recursos para financiar suas atividades e como esses recursos foram utilizados.')

# Demonstração das Mutações do Patrimônio Líquido (DMPL)
doc.add_heading('Demonstração das Mutações do Patrimônio Líquido (DMPL)', level=2)
doc.add_paragraph('A DMPL demonstra a evolução (aumento ou redução) do patrimônio líquido da entidade durante um período.\n'
                  'Ela é obrigatória para empresas estatais dependentes constituídas sob a forma de sociedades anônimas e facultativa para os demais órgãos e entidades dos entes da Federação.\n'
                  'Apresenta o déficit ou superávit patrimonial do período, cada mutação no patrimônio líquido reconhecida diretamente no mesmo, o efeito decorrente da mudança nos critérios contábeis e os efeitos decorrentes da retificação de erros cometidos em exercícios anteriores.')

# Notas Explicativas
doc.add_heading('Notas Explicativas', level=2)
doc.add_paragraph('As notas explicativas são informações adicionais às apresentadas nos quadros das DCASP e são consideradas parte integrante das demonstrações contábeis.\n'
                  'Elas têm como objetivo facilitar a compreensão das demonstrações contábeis e englobam informações exigidas pela lei, pelas normas contábeis e outras informações relevantes não suficientemente evidenciadas ou que não constam nas demonstrações contábeis.')

# Custos no Setor Público
doc.add_heading('Custos no Setor Público', level=1)
doc.add_paragraph('A principal diferença em matéria de custos entre o setor público e privado é na geração da informação. O setor público apresenta uma interdependência da execução orçamentária.\n'
                  'Por motivos de simplificação, o setor público considera como custo os gastos desembolsáveis fixados na dotação orçamentária.\n'
                  'Por exemplo, o pagamento da despesa com pessoal gera um lançamento no grupo de contas de controle de custos, enquanto a despesa com amortização e/ou provisão transita apenas pelo grupo de contas patrimoniais, não ocorrendo registro no grupo de contas de controle de custos.')

# MCASP Capítulo 3: Receita Orçamentária
doc.add_heading('MCASP Capítulo 3: Receita Orçamentária', level=1)
doc.add_paragraph('A receita orçamentária é o ingresso de recursos financeiros que passam a integrar o patrimônio público. As principais etapas da receita orçamentária são:')
doc.add_paragraph('a. Previsão\nb. Lançamento\nc. Arrecadação\nd. Recolhimento')

# MCASP Capítulo 4: Despesa Orçamentária
doc.add_heading('MCASP Capítulo 4: Despesa Orçamentária', level=1)
doc.add_paragraph('A despesa orçamentária é a aplicação de recursos financeiros para a execução de políticas públicas. As etapas da despesa orçamentária são:')
doc.add_paragraph('a. Planejamento\nb. Execução\nc. Liquidação\nd. Pagamento')

# Lançamentos Contábeis Públicos
doc.add_heading('Lançamentos Contábeis Públicos', level=1)
doc.add_paragraph('Os lançamentos contábeis públicos são divididos em três categorias principais: patrimonial, orçamentário e de controle.')
doc.add_heading('Lançamentos Patrimoniais', level=2)
doc.add_paragraph('Os lançamentos patrimoniais registram as variações nos ativos, passivos e patrimônio líquido das entidades públicas.')
doc.add_paragraph('Exemplo: Compra de um veículo\n'
                  'D - Veículos (Ativo Imobilizado)\n'
                  'C - Caixa (Ativo Circulante)')
doc.add_heading('Lançamentos Orçamentários', level=2)
doc.add_paragraph('Os lançamentos orçamentários registram a previsão e a execução do orçamento público, tanto das receitas quanto das despesas.')
doc.add_paragraph('Exemplo: Empenho de uma despesa\n'
                  'D - Despesa Empenhada (Orçamentário)\n'
                  'C - Crédito Disponível (Orçamentário)')
doc.add_heading('Lançamentos de Controle', level=2)
doc.add_paragraph('Os lançamentos de controle registram os atos potenciais que podem ou não afetar o patrimônio da entidade.')
doc.add_paragraph('Exemplo: Registro de garantias contratuais\n'
                  'D - Garantias Recebidas (Controle)\n'
                  'C - Garantias Concedidas (Controle)')

# Save the document
file_path = 'Resumo_Estudo_Prova_Final.docx'
doc.save(file_path)

print("Resumo salvo como 'Resumo_Estudo_Prova_Final.docx'")
